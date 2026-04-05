"""
Document parser — extracts text from uploaded PDFs, DOCX, and text files.

Uses PyMuPDF for PDFs (fast, reliable) and python-docx for Word documents.
Falls back to plain text reading for .txt, .md, .csv.
"""

import os
from pathlib import Path

import structlog
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json"}


def parse_file(file_path: str | Path) -> str:
    """
    Extract text content from a file.

    Args:
        file_path: Path to the uploaded file.

    Returns:
        Extracted text content.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    logger.info("Parsing file", path=str(path), ext=ext, size=path.stat().st_size)

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(path)
    else:
        return _parse_text(path)


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(str(path))
    pages = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"[Page {page_num + 1}]\n{text}")
    doc.close()
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    doc = DocxDocument(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _parse_text(path: Path) -> str:
    """Read plain text files."""
    return path.read_text(encoding="utf-8", errors="replace")


def parse_multiple(file_paths: list[str | Path]) -> list[str]:
    """Parse multiple files, returning a list of extracted texts."""
    results = []
    for fp in file_paths:
        try:
            results.append(parse_file(fp))
        except Exception as e:
            logger.error("Failed to parse file", path=str(fp), error=str(e))
            results.append(f"[Error parsing {Path(fp).name}: {e}]")
    return results
